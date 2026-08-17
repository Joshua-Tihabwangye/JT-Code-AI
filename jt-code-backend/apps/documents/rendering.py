from __future__ import annotations

import io

import markdown
from django.conf import settings

from apps.documents.models import Document


def _render_html(content: str) -> str:
    body = markdown.markdown(content or '', extensions=['tables', 'fenced_code', 'sane_lists'])
    return f"""<!DOCTYPE html>
<html>
<head>
<meta charset="utf-8">
<style>
  body {{ font-family: 'DejaVu Sans', sans-serif; margin: 40px; color: #1a1a1a; line-height: 1.6; }}
  h1, h2, h3 {{ color: #0f172a; }}
  table {{ border-collapse: collapse; width: 100%; margin: 12px 0; }}
  th, td {{ border: 1px solid #cbd5e1; padding: 6px 10px; text-align: left; }}
  pre {{ background: #f1f5f9; padding: 12px; border-radius: 6px; overflow-x: auto; }}
  code {{ font-family: 'DejaVu Sans Mono', monospace; }}
  blockquote {{ border-left: 4px solid #94a3b8; margin-left: 0; padding-left: 16px; color: #475569; }}
</style>
</head>
<body>{body}</body>
</html>"""


def render_pdf(document: Document) -> bytes:
    from weasyprint import HTML

    html = _render_html(document.content)
    pdf = HTML(string=html).write_pdf()
    return pdf


def render_docx(document: Document) -> bytes:
    from docx import Document as DocxDocument

    buffer = io.BytesIO()
    doc = DocxDocument()
    doc.add_heading(document.title, level=0)
    for block in _split_blocks(document.content):
        if block.startswith('# '):
            doc.add_heading(block[2:].strip(), level=1)
        elif block.startswith('## '):
            doc.add_heading(block[3:].strip(), level=2)
        elif block.startswith('- '):
            doc.add_paragraph(block[2:].strip(), style='List Bullet')
        else:
            doc.add_paragraph(block.strip())
    doc.save(buffer)
    return buffer.getvalue()


def _split_blocks(content: str) -> list[str]:
    return [b for b in content.split('\n\n') if b.strip()]


def upload_bytes_to_cloudinary(content: bytes, public_id: str, resource_type: str = 'raw') -> str | None:
    if not all((settings.CLOUDINARY_CLOUD_NAME, settings.CLOUDINARY_API_KEY, settings.CLOUDINARY_API_SECRET)):
        return None
    if settings.CLOUDINARY_CLOUD_NAME == 'replace_me':
        return None
    try:
        import cloudinary.uploader

        result = cloudinary.uploader.upload(
            io.BytesIO(content),
            public_id=public_id,
            resource_type=resource_type,
            overwrite=True,
        )
        return result.get('secure_url')
    except Exception:
        return None