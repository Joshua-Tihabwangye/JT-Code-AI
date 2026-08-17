from __future__ import annotations

import io
import shutil
import subprocess
from pathlib import Path

import markdown
from django.conf import settings

from apps.conversions.models import ConversionJob
from apps.conversions.serializers import ALLOWED_MATRIX
from apps.documents.rendering import upload_bytes_to_cloudinary

CONVERSION_ROOT = Path(settings.BASE_DIR) / 'converted_files'


def _html_from_markdown(content: str) -> str:
    body = markdown.markdown(content or '', extensions=['tables', 'fenced_code', 'sane_lists'])
    return f"""<!DOCTYPE html>
<html>
<head>
<meta charset="utf-8">
<style>
  body {{ font-family: 'DejaVu Sans', sans-serif; margin: 40px; color: #1a1a1a; line-height: 1.6; }}
  h1, h2, h3 {{ color: #0f172a; }}
  table {{ border-collapse: collapse; width: 100%; margin: 12px 0; }}
  th, td {{ border: 1px solid #cbd5e1; padding: 6px 10px; text-align: left; }}
  pre {{ background: #f1f5f9; padding: 12px; border-radius: 6px; overflow-x: auto; }}
  code {{ font-family: 'DejaVu Sans Mono', monospace; }}
</style>
</head>
<body>{body}</body>
</html>"""


def _pdf_from_html(html: str) -> bytes:
    from weasyprint import HTML

    return HTML(string=html).write_pdf()


def _docx_from_markdown(content: str) -> bytes:
    from docx import Document as DocxDocument

    buffer = io.BytesIO()
    doc = DocxDocument()
    for block in [b for b in (content or '').split('\n\n') if b.strip()]:
        if block.startswith('# '):
            doc.add_heading(block[2:].strip(), level=1)
        elif block.startswith('## '):
            doc.add_heading(block[3:].strip(), level=2)
        elif block.startswith('- '):
            doc.add_paragraph(block[2:].strip(), style='List Bullet')
        else:
            doc.add_paragraph(block.strip())
    doc.save(buffer)
    return buffer.getvalue()


def _text_from_pdf(content: bytes) -> tuple[str, str]:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(content))
    pages = [page.extract_text() or '' for page in reader.pages]
    text = '\n\n'.join(pages)
    return text, 'txt'


def _markdown_from_pdf(content: bytes) -> tuple[str, str]:
    text, _ = _text_from_pdf(content)
    return text, 'md'


def _image_convert(content: bytes, output_format: str) -> bytes:
    from PIL import Image

    image = Image.open(io.BytesIO(content))
    if image.mode not in ('RGB', 'RGBA'):
        image = image.convert('RGB')
    buffer = io.BytesIO()
    save_format = 'JPEG' if output_format == 'jpg' else output_format.upper()
    image.save(buffer, format=save_format)
    return buffer.getvalue()


def _docx_to_pdf(input_path: Path, output_path: Path) -> None:
    liboffice = (
        shutil.which('libreoffice')
        or shutil.which('soffice')
        or '/snap/bin/libreoffice'
    )
    if not Path(liboffice).exists():
        raise RuntimeError('LibreOffice is not installed; DOCX->PDF conversion unavailable.')
    subprocess.run(
        [
            liboffice, '--headless', '--convert-to', 'pdf',
            '--outdir', str(output_path.parent), str(input_path),
        ],
        check=True,
        capture_output=True,
        timeout=120,
    )
    generated = output_path.parent / f'{output_path.stem}.pdf'
    if generated.exists():
        generated.rename(output_path)


def run_conversion(job: ConversionJob) -> bytes:
    if (job.input_format, job.output_format) not in ALLOWED_MATRIX:
        raise ValueError(
            f'Conversion {job.input_format}->{job.output_format} is not allowed.'
        )

    content = Path(job.input_path).read_bytes() if job.input_path else b''

    if job.input_format == 'pdf' and job.output_format == 'txt':
        output, _ = _text_from_pdf(content)
        return output.encode('utf-8')
    if job.input_format == 'pdf' and job.output_format == 'md':
        output, _ = _markdown_from_pdf(content)
        return output.encode('utf-8')
    if job.input_format in {'md', 'txt'} and job.output_format == 'pdf':
        return _pdf_from_html(_html_from_markdown(content.decode('utf-8', errors='replace')))
    if job.input_format == 'html' and job.output_format == 'pdf':
        return _pdf_from_html(content.decode('utf-8', errors='replace'))
    if job.input_format in {'md', 'txt'} and job.output_format == 'docx':
        return _docx_from_markdown(content.decode('utf-8', errors='replace'))
    if job.input_format in {'png', 'jpg', 'jpeg', 'webp'}:
        return _image_convert(content, job.output_format)
    if job.input_format == 'docx' and job.output_format == 'pdf':
        if not job.input_path:
            raise ValueError('DOCX->PDF conversion requires an uploaded file.')
        input_path = Path(job.input_path)
        output_path = CONVERSION_ROOT / f'{job.id}.pdf'
        output_path.parent.mkdir(parents=True, exist_ok=True)
        _docx_to_pdf(input_path, output_path)
        return output_path.read_bytes()
    raise ValueError(f'Conversion {job.input_format}->{job.output_format} is not implemented.')


def finalize_conversion(job: ConversionJob, output: bytes) -> str | None:
    CONVERSION_ROOT.mkdir(parents=True, exist_ok=True)
    local_path = CONVERSION_ROOT / f'{job.id}.{job.output_format}'
    local_path.write_bytes(output)
    job.output_path = str(local_path)
    job.output_bytes = len(output)
    cloudinary_url = upload_bytes_to_cloudinary(
        output,
        public_id=f'conversions/{job.owner_id}/{job.id}',
    )
    if cloudinary_url:
        job.output_url = cloudinary_url
    return cloudinary_url
